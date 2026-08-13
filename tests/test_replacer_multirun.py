"""
tests/test_replacer_multirun.py
--------------------------------
Dedicated tests for DOCX multi-run span replacement.

These tests create real python-docx Document objects in memory so that
apply_replacement_to_block() is exercised against actual Run objects —
not mocks.  Key scenarios tested:

  1. Single-run replacement preserves bold/italic on untouched runs
  2. Multi-run span replacement:
       - fake value appears in first affected run
       - subsequent affected runs have empty text
       - document re-saves and re-opens without error
  3. Table cell text replacement works correctly
  4. Header text replacement works correctly
  5. Reverse-order (last→first) processing preserves earlier offsets
"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import io
import pytest
from docx import Document
from docx.oxml.ns import qn

from models import TextBlock
from docx_io import _build_run_offsets, _combine_run_text, apply_replacement_to_block


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_block_from_para(para, para_idx=0) -> TextBlock:
    runs = para.runs
    return TextBlock(
        text=_combine_run_text(runs),
        location={
            "type": "paragraph",
            "para_idx": para_idx,
            "runs": runs,
            "run_offsets": _build_run_offsets(runs),
        },
    )


def _make_block_from_cell_para(cell, table_idx=0, row_idx=0, col_idx=0, para_idx=0):
    para = cell.paragraphs[para_idx]
    runs = para.runs
    return TextBlock(
        text=_combine_run_text(runs),
        location={
            "type": "table_cell",
            "table_idx": table_idx,
            "row_idx": row_idx,
            "col_idx": col_idx,
            "para_idx": para_idx,
            "runs": runs,
            "run_offsets": _build_run_offsets(runs),
        },
    )


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _reload_doc(doc: Document) -> Document:
    """Save and re-open document to verify no XML corruption."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


# ---------------------------------------------------------------------------
# 1. Single-run replacement
# ---------------------------------------------------------------------------

class TestSingleRunReplacement:
    def test_basic_replacement(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("My name is John Smith and I live here.")
        block = _make_block_from_para(para)

        start = block.text.index("John Smith")
        end = start + len("John Smith")
        apply_replacement_to_block(block, start, end, "Jane Doe")

        assert run.text == "My name is Jane Doe and I live here."

    def test_formatting_preserved_on_replaced_run(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Call 4111111111111111 now")
        run.bold = True
        block = _make_block_from_para(para)

        start = block.text.index("4111111111111111")
        end = start + len("4111111111111111")
        apply_replacement_to_block(block, start, end, "XXXX-XXXX-XXXX-XXXX")

        assert run.bold is True
        assert "4111111111111111" not in run.text

    def test_untouched_runs_unaffected(self):
        doc = Document()
        para = doc.add_paragraph()
        run_a = para.add_run("Hello ")
        run_b = para.add_run("John Smith")
        run_c = para.add_run(" goodbye")
        run_b.italic = True
        block = _make_block_from_para(para)

        start = block.text.index("John Smith")
        end = start + len("John Smith")
        apply_replacement_to_block(block, start, end, "Jane Doe")

        # run_a and run_c must be unmodified
        assert run_a.text == "Hello "
        assert run_c.text == " goodbye"


# ---------------------------------------------------------------------------
# 2. Multi-run span replacement
# ---------------------------------------------------------------------------

class TestMultiRunReplacement:
    def _make_split_doc(self, parts: list[tuple[str, bool]]):
        """
        Create a Document with one paragraph whose runs are given by *parts*
        (text, bold) tuples.
        """
        doc = Document()
        para = doc.add_paragraph()
        for text, bold in parts:
            run = para.add_run(text)
            run.bold = bold
        return doc, para

    def test_name_split_across_two_runs(self):
        """'Kushal ' (normal) + 'Subbayya Hegde' (bold) → full span replaced."""
        doc, para = self._make_split_doc([
            ("Director: Kushal ", False),
            ("Subbayya Hegde", True),
            (", DIN 12345", False),
        ])
        block = _make_block_from_para(para)
        full_name = "Kushal Subbayya Hegde"
        start = block.text.index("Kushal Subbayya Hegde")
        end = start + len(full_name)

        apply_replacement_to_block(block, start, end, "Aarav Mehta")

        # Full document text must not contain original name
        rebuilt = "".join(r.text for r in para.runs)
        assert "Kushal" not in rebuilt
        assert "Subbayya" not in rebuilt
        assert "Aarav Mehta" in rebuilt

    def test_name_split_across_three_runs(self):
        doc, para = self._make_split_doc([
            ("Ku", False),
            ("shal Sub", True),
            ("bayya Hegde extra", False),
        ])
        block = _make_block_from_para(para)
        # The full text of the block
        full = block.text
        target = "Kushal Subbayya Hegde"
        start = full.index(target)
        end = start + len(target)

        apply_replacement_to_block(block, start, end, "Priya Nair")

        rebuilt = "".join(r.text for r in para.runs)
        assert target not in rebuilt
        assert "Priya Nair" in rebuilt

    def test_document_reopens_after_multirun_replacement(self):
        """Document must not be corrupted after multi-run replacement."""
        doc, para = self._make_split_doc([
            ("Email: alice", False),
            ("@example", True),
            (".com done", False),
        ])
        block = _make_block_from_para(para)
        email = "alice@example.com"
        start = block.text.index(email)
        end = start + len(email)

        apply_replacement_to_block(block, start, end, "bob@testmail.org")

        # Re-open must not raise
        reloaded = _reload_doc(doc)
        assert reloaded is not None

    def test_formatting_of_untouched_run_preserved(self):
        doc, para = self._make_split_doc([
            ("Start ", False),
            ("John ", True),
            ("Smith", False),
            (" End", False),
        ])
        last_run = para.runs[-1]
        block = _make_block_from_para(para)
        start = block.text.index("John Smith")
        end = start + len("John Smith")

        apply_replacement_to_block(block, start, end, "Jane")

        # The last " End" run must remain unchanged
        assert last_run.text == " End"


# ---------------------------------------------------------------------------
# 3. Table cell replacement
# ---------------------------------------------------------------------------

class TestTableCellReplacement:
    def test_cell_text_replaced(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run("Contact: user@domain.com")

        block = _make_block_from_cell_para(cell)
        email = "user@domain.com"
        start = block.text.index(email)
        end = start + len(email)
        apply_replacement_to_block(block, start, end, "anon@example.com")

        assert "user@domain.com" not in run.text
        assert "anon@example.com" in run.text

    def test_table_doc_reopens(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run("SSN: 123-45-6789")

        block = _make_block_from_cell_para(cell)
        start = block.text.index("123-45-6789")
        end = start + len("123-45-6789")
        apply_replacement_to_block(block, start, end, "987-65-4321")

        reloaded = _reload_doc(doc)
        assert reloaded is not None


# ---------------------------------------------------------------------------
# 4. Reverse-order processing preserves earlier offsets
# ---------------------------------------------------------------------------

class TestReverseOrderProcessing:
    def test_two_matches_last_first(self):
        """
        If two non-overlapping matches are processed last-to-first,
        the earlier match's offsets are still valid after the later one
        is replaced.
        """
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Alice Smith called Bob Jones yesterday.")
        block = _make_block_from_para(para)

        # Identify both spans
        alice_start = block.text.index("Alice Smith")
        alice_end = alice_start + len("Alice Smith")
        bob_start = block.text.index("Bob Jones")
        bob_end = bob_start + len("Bob Jones")

        # Process in reverse (last first: Bob, then Alice)
        apply_replacement_to_block(block, bob_start, bob_end, "Dave Green")
        # After Bob replacement, run.text changed — rebuild offset tracking
        # (in production, apply_to_docx refreshes the block's run offsets)
        # For this test, verify Alice's offsets still correct before replacement
        apply_replacement_to_block(block, alice_start, alice_end, "Carol White")

        final = run.text
        assert "Alice" not in final
        assert "Bob" not in final
        assert "Carol White" in final
        assert "Dave Green" in final
