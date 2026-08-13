"""
tests/test_docx_io.py
----------------------
Tests for DOCX text extraction and run-offset mapping (docx_io.py).

Uses in-memory Document objects (no fixture files needed).
"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import io
import pytest
from docx import Document

from docx_io import (
    load_document,
    extract_all_text_blocks,
    iter_paragraph_blocks,
    iter_table_cell_blocks,
    iter_header_footer_blocks,
    _build_run_offsets,
    _combine_run_text,
    save_document,
)


def _save_and_path(doc: Document, tmp_path) -> str:
    p = tmp_path / "test.docx"
    doc.save(str(p))
    return str(p)


# ---------------------------------------------------------------------------
# _build_run_offsets
# ---------------------------------------------------------------------------

class TestRunOffsets:
    def test_empty_runs(self):
        assert _build_run_offsets([]) == []

    def test_single_run(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello")
        assert _build_run_offsets(para.runs) == [0]

    def test_multiple_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")   # len 6
        para.add_run("world")   # len 5
        para.add_run("!")        # len 1
        offsets = _build_run_offsets(para.runs)
        assert offsets == [0, 6, 11]

    def test_offset_plus_runlen_equals_next_offset(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("AB")
        para.add_run("CDE")
        para.add_run("F")
        offsets = _build_run_offsets(para.runs)
        runs = para.runs
        for i in range(len(runs) - 1):
            assert offsets[i] + len(runs[i].text) == offsets[i + 1]


# ---------------------------------------------------------------------------
# iter_paragraph_blocks
# ---------------------------------------------------------------------------

class TestParagraphBlocks:
    def test_extracts_body_paragraphs(self, tmp_path):
        doc = Document()
        doc.add_paragraph("First paragraph text.")
        doc.add_paragraph("Second paragraph text.")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_paragraph_blocks(loaded))
        texts = [b.text for b in blocks]
        assert any("First paragraph" in t for t in texts)
        assert any("Second paragraph" in t for t in texts)

    def test_empty_paragraphs_skipped(self, tmp_path):
        doc = Document()
        doc.add_paragraph("")        # empty — should be skipped
        doc.add_paragraph("   ")     # whitespace-only — skipped
        doc.add_paragraph("Real content here.")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_paragraph_blocks(loaded))
        assert all(b.text.strip() for b in blocks)

    def test_location_has_para_idx(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Para zero.")
        doc.add_paragraph("Para one.")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_paragraph_blocks(loaded))
        # All blocks must have para_idx in location
        assert all("para_idx" in b.location for b in blocks)

    def test_run_offsets_present(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        para.add_run("Hello ")
        para.add_run("world")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_paragraph_blocks(loaded))
        real_blocks = [b for b in blocks if "Hello" in b.text]
        assert real_blocks, "Expected block with 'Hello world' text"
        b = real_blocks[0]
        assert "run_offsets" in b.location
        assert b.location["run_offsets"][0] == 0


# ---------------------------------------------------------------------------
# iter_table_cell_blocks
# ---------------------------------------------------------------------------

class TestTableCellBlocks:
    def test_extracts_table_text(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Email"
        table.rows[1].cells[0].text = "John Smith"
        table.rows[1].cells[1].text = "john@example.com"
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_table_cell_blocks(loaded))
        texts = [b.text for b in blocks]
        assert any("John Smith" in t for t in texts)
        assert any("john@example.com" in t for t in texts)

    def test_location_has_table_row_col_idx(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Cell A"
        table.rows[0].cells[1].text = "Cell B"
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_table_cell_blocks(loaded))
        for b in blocks:
            assert "table_idx" in b.location
            assert "row_idx" in b.location
            assert "col_idx" in b.location


# ---------------------------------------------------------------------------
# iter_header_footer_blocks
# ---------------------------------------------------------------------------

class TestHeaderFooterBlocks:
    def test_extracts_header_text(self, tmp_path):
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "CONFIDENTIAL — Header PII: john@example.com"
        doc.add_paragraph("Body text.")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_header_footer_blocks(loaded))
        texts = [b.text for b in blocks]
        assert any("CONFIDENTIAL" in t or "john@example.com" in t for t in texts)

    def test_location_type_is_header_or_footer(self, tmp_path):
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Header content"
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = list(iter_header_footer_blocks(loaded))
        if blocks:
            assert all(b.location["type"] in ("header", "footer") for b in blocks)


# ---------------------------------------------------------------------------
# extract_all_text_blocks
# ---------------------------------------------------------------------------

class TestExtractAllTextBlocks:
    def test_non_empty_result(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Some content here.")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = extract_all_text_blocks(loaded)
        assert len(blocks) > 0

    def test_paragraphs_and_table_cells_combined(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Body paragraph.")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Table cell content."
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = extract_all_text_blocks(loaded)
        texts = [b.text for b in blocks]
        assert any("Body paragraph" in t for t in texts)
        assert any("Table cell content" in t for t in texts)

    def test_character_offset_mapping(self, tmp_path):
        """run_offsets must allow locating a substring within block.text."""
        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        para.add_run("Hello ")
        para.add_run("John Smith")
        para.add_run(" goodbye")
        path = _save_and_path(doc, tmp_path)

        loaded = load_document(path)
        blocks = extract_all_text_blocks(loaded)
        target_block = next((b for b in blocks if "John Smith" in b.text), None)
        assert target_block is not None

        start = target_block.text.index("John Smith")
        end = start + len("John Smith")
        offsets = target_block.location["run_offsets"]
        runs = target_block.location["runs"]

        # Find which run contains the name
        containing_run = None
        for i, off in enumerate(offsets):
            run_end = offsets[i + 1] if i + 1 < len(offsets) else off + len(runs[i].text)
            if off <= start < run_end:
                containing_run = runs[i]
                break

        assert containing_run is not None
        assert "John Smith" in containing_run.text
