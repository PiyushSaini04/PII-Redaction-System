"""
tests/test_ner_validation.py
-----------------------------
Regression tests for the precision-first NER validation layer.

Covers:
- Every user-reported false positive (must NOT be detected)
- Every traced fabrication original span (must NOT be detected)
- All confirmed true positives (MUST be detected with correct category)
- Exact character-span verification for key entities
- DOCX replacement isolation (only PII span changes)
- Category-conflict resolution (company name via PERSON label)
"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from models import TextBlock, PIIMatch
from detectors.ner_detectors import (
    detect_person_names,
    detect_organizations,
    _validate_person,
    _validate_company,
    _clean_entity_text,
)


def block(text):
    return TextBlock(
        text=text,
        location={"type": "paragraph", "para_idx": 0, "runs": [], "run_offsets": []},
    )


# ---------------------------------------------------------------------------
# _clean_entity_text
# ---------------------------------------------------------------------------

class TestCleanEntityText:
    def test_strips_trailing_ampersand(self):
        assert _clean_entity_text("Dhaulagiri Family Trust&") == "Dhaulagiri Family Trust"

    def test_strips_footnote_markers(self):
        assert _clean_entity_text("Kushal Subbayya Hegde*^&") == "Kushal Subbayya Hegde"

    def test_strips_hash(self):
        assert _clean_entity_text("Services Tax#") == "Services Tax"

    def test_leaves_clean_text_unchanged(self):
        assert _clean_entity_text("KSH International Limited") == "KSH International Limited"

    def test_leaves_hyphen_inside(self):
        assert _clean_entity_text("Kirtane & Pandit LLP") == "Kirtane & Pandit LLP"


# ---------------------------------------------------------------------------
# _validate_person — traced fabrication originals (must all be REJECTED)
# ---------------------------------------------------------------------------

class TestValidatePersonTracedFPs:
    """
    These are the ORIGINAL texts that produced fabricated replacements
    in the redacted DOCX output. Each must be rejected by the validator.
    """
    def test_non_gaap_measures_rejected_as_person(self):
        # Was replaced by 'Walter, Edwards and Rios' — originated as COMPANY FP
        # but verify it is also rejected from PERSON
        r = _validate_person("Non-GAAP Measures", "Non-GAAP Measures", "")
        assert not r.accepted, "Non-GAAP Measures must not be PERSON"

    def test_broad_family_trust_not_person(self):
        # Was replaced by 'Arroyo, Miller and Tucker'
        r = _validate_person("Broad Family Trust", "Broad Family Trust", "")
        assert not r.accepted or r.suggested_category == "COMPANY", \
            "Broad Family Trust must not be PERSON"

    def test_makalu_family_trust_not_person(self):
        # Was replaced by 'Allen-Allen'
        r = _validate_person("Makalu Family Trust", "Makalu Family Trust", "")
        assert not r.accepted or r.suggested_category == "COMPANY"

    def test_annapurna_family_trust_not_person(self):
        # Was replaced by 'Spence PLC'
        r = _validate_person("Annapurna Family Trust", "Annapurna Family Trust", "")
        assert not r.accepted or r.suggested_category == "COMPANY"

    def test_post_offer_not_person(self):
        r = _validate_person("post-Offer", "post-Offer", "")
        assert not r.accepted

    def test_rohit_branch_not_person(self):
        r = _validate_person("Rohit Branch", "Rohit Branch", "")
        assert not r.accepted


# ---------------------------------------------------------------------------
# _validate_company — traced fabrication originals (must be REJECTED)
# ---------------------------------------------------------------------------

class TestValidateCompanyTracedFPs:
    def test_non_gaap_measures_rejected(self):
        r = _validate_company("Non-GAAP Measures", "Non-GAAP Measures", "")
        assert not r.accepted, f"Expected reject, got reason={r.reason}"

    def test_broad_family_trust_rejected_per_policy(self):
        """User policy: promoter trusts are NOT redacted."""
        r = _validate_company("Broad Family Trust", "Broad Family Trust", "")
        assert not r.accepted

    def test_securities_and_exchange_board_of_india_rejected(self):
        r = _validate_company("Securities and Exchange Board of India", "Securities and Exchange Board of India", "")
        assert not r.accepted, "Must be rejected per user instruction"
        
    def test_securities_act_rejected(self):
        r = _validate_company("Securities and Exchange Board of India Act", "Securities and Exchange Board of India Act", "")
        assert not r.accepted, "Acts must be rejected"

    def test_post_offer_rejected(self):
        r = _validate_company("post-Offer", "post-Offer", "")
        assert not r.accepted
        r2 = _validate_company("post-Offer Equity Share", "post-Offer Equity Share", "")
        assert not r2.accepted

    def test_offer_for_rejected(self):
        r = _validate_company("OFFER FOR", "OFFER FOR", "")
        assert not r.accepted

    def test_rohit_branch_rejected(self):
        r = _validate_company("Rohit Branch", "Rohit Branch", "")
        assert not r.accepted

    def test_makalu_family_trust_rejected_per_policy(self):
        r = _validate_company("Makalu Family Trust", "Makalu Family Trust", "")
        assert not r.accepted

    def test_annapurna_family_trust_rejected_per_policy(self):
        r = _validate_company("Annapurna Family Trust", "Annapurna Family Trust", "")
        assert not r.accepted

    def test_dhaulagiri_family_trust_rejected_per_policy(self):
        r = _validate_company("DHAULAGIRI FAMILY TRUST", "DHAULAGIRI FAMILY TRUST", "")
        assert not r.accepted

    def test_freehold_land_rejected(self):
        r = _validate_company("Freehold Land and Leasehold Land", "", "")
        assert not r.accepted

    def test_bank_balances_rejected(self):
        r = _validate_company("Bank Balances and Advances", "", "")
        assert not r.accepted

    def test_employer_contribution_rejected(self):
        r = _validate_company("Employer and Employee Contribution", "", "")
        assert not r.accepted

    def test_tower_2a_2b_rejected(self):
        r = _validate_company("Tower 2A & 2B", "", "")
        assert not r.accepted

    def test_sebi_regulations_rejected(self):
        r = _validate_company(
            "Securities and Exchange Board of India (Foreign Venture Capital Investor) Regulations",
            "", "")
        assert not r.accepted, "Regulation name should be rejected via act_regulation_name"

    def test_callings_employment_act_rejected(self):
        r = _validate_company("Callings and Employment Act", "", "")
        assert not r.accepted

    def test_services_tax_ampersand_rejected(self):
        # Cleaned: "Services Tax" — no positive evidence
        r = _validate_company("Services Tax&", "", "")
        assert not r.accepted

    def test_kushal_hegde_with_markers_rejected(self):
        # Person name with footnote markers appearing as ORG
        r = _validate_company("Kushal Subbayya Hegde*^&", "", "")
        assert not r.accepted

    def test_rakhi_girija_shetty_rejected(self):
        r = _validate_company("Rakhi Girija Shetty", "", "")
        assert not r.accepted


# ---------------------------------------------------------------------------
# _validate_person — previously reported FPs (must all STAY rejected)
# ---------------------------------------------------------------------------

class TestValidatePersonPreviousFPs:
    def test_rejects_offer(self):
        assert not _validate_person("Offer", "", "").accepted

    def test_rejects_promoters(self):
        assert not _validate_person("Promoters", "", "").accepted

    def test_rejects_directors(self):
        assert not _validate_person("Directors", "", "").accepted

    def test_rejects_cap_price(self):
        assert not _validate_person("Cap Price", "", "").accepted

    def test_rejects_floor_price(self):
        assert not _validate_person("Floor Price", "", "").accepted

    def test_rejects_allcaps_brlm(self):
        assert not _validate_person("BRLM", "", "").accepted

    def test_rejects_allcaps_risks(self):
        assert not _validate_person("RISKS", "", "").accepted

    def test_rejects_iso_code(self):
        assert not _validate_person("ISO 9001:2015", "", "").accepted

    def test_rejects_room_code(self):
        assert not _validate_person("C-101", "", "").accepted

    def test_rejects_branch_suffix(self):
        assert not _validate_person("Parents Branch", "", "").accepted
        assert not _validate_person("Rajesh Branch", "", "").accepted

    def test_rejects_share_transfer_agents(self):
        assert not _validate_person("Share Transfer Agents", "", "").accepted

    def test_rejects_location(self):
        assert not _validate_person("Deccan Gymkhana", "", "").accepted
        assert not _validate_person("Chakan Taluka - Khed", "", "").accepted

    def test_rejects_pat_cagr(self):
        assert not _validate_person("PAT CAGR", "", "").accepted

    def test_rejects_section_heading(self):
        assert not _validate_person("B.  Non-GAAP Measures", "", "").accepted

    def test_rejects_bid_singular(self):
        assert not _validate_person("Bid", "", "").accepted

    def test_rejects_nuvama(self):
        assert not _validate_person("Nuvama", "", "").accepted

    def test_rejects_vikhroli(self):
        assert not _validate_person("Vikhroli", "", "").accepted


# ---------------------------------------------------------------------------
# _validate_company — previously reported FPs (must all STAY rejected)
# ---------------------------------------------------------------------------

class TestValidateCompanyPreviousFPs:
    def test_rejects_definitions(self):
        assert not _validate_company("DEFINITIONS", "DEFINITIONS", "").accepted

    def test_rejects_currency(self):
        assert not _validate_company("CURRENCY", "CURRENCY", "").accepted

    def test_rejects_board(self):
        assert not _validate_company("Board", "", "").accepted

    def test_rejects_board_of_directors(self):
        assert not _validate_company("Board of Directors", "", "").accepted

    def test_rejects_the_offer_price(self):
        assert not _validate_company("the Offer Price", "", "").accepted

    def test_rejects_offer_for_sale(self):
        assert not _validate_company("the Offer for Sale", "", "").accepted

    def test_rejects_ipo_allotment(self):
        assert not _validate_company("Allotment", "", "").accepted

    def test_rejects_ipo_bonus(self):
        assert not _validate_company("Bonus", "", "").accepted

    def test_rejects_clause_fragment(self):
        assert not _validate_company("THE OFFER SHALL CONSTITUTE", "", "").accepted

    def test_rejects_equity_share_capital(self):
        assert not _validate_company("EQUITY SHARE CAPITAL OF OUR COMPANY", "", "").accepted

    def test_rejects_red_herring_prospectus(self):
        assert not _validate_company("RED HERRING PROSPECTUS", "", "").accepted
        assert not _validate_company("Red Herring Prospectus", "", "").accepted


# ---------------------------------------------------------------------------
# True positives — T1 company suffix (must be ACCEPTED)
# ---------------------------------------------------------------------------

class TestValidateCompanyTruePositives:
    def test_accepts_ksh_international_limited(self):
        r = _validate_company("KSH INTERNATIONAL LIMITED", "", "")
        assert r.accepted
        assert r.suggested_category == "COMPANY"
        assert r.reason == "company_suffix"

    def test_accepts_ksh_international_mixed(self):
        r = _validate_company("KSH International Limited", "", "")
        assert r.accepted

    def test_accepts_bhandary_metal(self):
        r = _validate_company("Bhandary Metal Extrusion Private Limited", "", "")
        assert r.accepted
        assert r.reason == "company_suffix"

    def test_accepts_kirtane_pandit_llp(self):
        r = _validate_company("Kirtane & Pandit LLP", "", "")
        assert r.accepted
        assert r.reason == "company_suffix"

    def test_accepts_hdfc_bank_limited(self):
        r = _validate_company("HDFC Bank Limited", "", "")
        assert r.accepted

    def test_accepts_waterloo_industrial_as_company(self):
        """Company name redirected from PERSON label via company suffix."""
        r = _validate_person("Waterloo Industrial Park VI Private Limited", "", "")
        assert r.accepted
        assert r.suggested_category == "COMPANY"


# ---------------------------------------------------------------------------
# True positives — T3 public body (must be ACCEPTED with needs_review)
# ---------------------------------------------------------------------------

class TestValidateCompanyPublicBodies:
    def test_accepts_rbi_abbreviation(self):
        r = _validate_company("RBI", "", "")
        assert r.accepted
        assert r.needs_review


# ---------------------------------------------------------------------------
# True positives — T5 person names (must be ACCEPTED)
# ---------------------------------------------------------------------------

class TestValidatePersonTruePositives:
    def test_accepts_sarthak_malvadkar(self):
        r = _validate_person("Sarthak Malvadkar", "", "")
        assert r.accepted
        assert r.suggested_category == "PERSON"

    def test_accepts_kushal_subbayya_hegde(self):
        assert _validate_person("Kushal Subbayya Hegde", "", "").accepted

    def test_accepts_pushpa_kushal_hegde(self):
        assert _validate_person("Pushpa Kushal Hegde", "", "").accepted

    def test_accepts_rajesh_kushal_hegde(self):
        assert _validate_person("Rajesh Kushal Hegde", "", "").accepted

    def test_accepts_sandesh_bhagwat(self):
        assert _validate_person("Sandesh Bhagwat", "", "").accepted

    def test_accepts_amod_joshi(self):
        assert _validate_person("Amod Joshi", "", "").accepted

    def test_accepts_shanti_gopalkrishnan(self):
        assert _validate_person("Shanti Gopalkrishnan", "", "").accepted

    def test_accepts_allcaps_name(self):
        """All-caps names appear in tables; must be accepted with needs_review."""
        r = _validate_person("KUSHAL SUBBAYYA HEGDE", "", "")
        assert r.accepted
        assert r.needs_review


# ---------------------------------------------------------------------------
# End-to-end detector regression
# ---------------------------------------------------------------------------

class TestPersonDetectorRegression:
    def test_offer_not_person(self):
        r = detect_person_names(block("The Offer opens on Monday."))
        assert not any(m.text == "Offer" for m in r)

    def test_kushal_hegde_detected(self):
        r = detect_person_names(block("Kushal Subbayya Hegde is a promoter of the Company."))
        assert any("Kushal" in m.text and "Hegde" in m.text for m in r)

    def test_non_gaap_not_person(self):
        r = detect_person_names(block("B.  Non-GAAP Measures and their reconciliation."))
        assert not any("Non-GAAP" in m.text for m in r)

    def test_company_suffix_entity_category_is_company(self):
        r = detect_person_names(block("Waterloo Industrial Park VI Private Limited holds shares."))
        assert not any(m.category == "PERSON" for m in r)


class TestCompanyDetectorRegression:
    def test_definitions_not_company(self):
        assert not detect_organizations(block("DEFINITIONS"))

    def test_currency_not_company(self):
        assert not detect_organizations(block("CURRENCY"))

    def test_board_not_company(self):
        r = detect_organizations(block("The resolution was passed by our Board."))
        assert not any(m.text == "Board" for m in r)

    def test_non_gaap_not_company(self):
        r = detect_organizations(block("Non-GAAP Measures are presented for convenience."))
        assert not any("Non-GAAP" in m.text for m in r)

    def test_ksh_international_detected(self):
        r = detect_organizations(block("The issuer is KSH International Limited."))
        assert any("KSH International" in m.text for m in r)

    def test_bhandary_metal_detected(self):
        r = detect_organizations(block("Subsidiary: Bhandary Metal Extrusion Private Limited."))
        assert any("Bhandary Metal Extrusion" in m.text for m in r)

    def test_family_trust_not_company_per_policy(self):
        r = detect_organizations(block("DHAULAGIRI FAMILY TRUST holds equity shares."))
        assert not any("FAMILY TRUST" in m.text for m in r)


# ---------------------------------------------------------------------------
# Exact span offset verification
# ---------------------------------------------------------------------------

class TestExactSpanOffset:
    def test_person_span_exact(self):
        text = "The director Kushal Subbayya Hegde signed the agreement."
        b = block(text)
        matches = detect_person_names(b)
        person = [m for m in matches if "Kushal" in m.text and "Hegde" in m.text]
        assert person, "Expected PERSON match for 'Kushal Subbayya Hegde'"
        m = person[0]
        assert text[m.start_char:m.end_char] == m.text
        assert "The director " not in m.text
        assert "signed the agreement" not in m.text

    def test_company_span_exact(self):
        text = "The issuer is KSH International Limited, incorporated in Pune."
        b = block(text)
        matches = detect_organizations(b)
        company = [m for m in matches if "KSH" in m.text]
        assert company, "Expected COMPANY match for 'KSH International Limited'"
        m = company[0]
        assert text[m.start_char:m.end_char] == m.text

    def test_non_pii_text_unchanged_after_replacement(self):
        from docx import Document
        from docx_io import _build_run_offsets, _combine_run_text, apply_replacement_to_block

        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        run = para.add_run("The director Kushal Subbayya Hegde signed the agreement.")

        runs = para.runs
        blk = TextBlock(
            text=_combine_run_text(runs),
            location={
                "type": "paragraph", "para_idx": 0,
                "runs": runs, "run_offsets": _build_run_offsets(runs),
            },
        )
        person = [m for m in detect_person_names(blk)
                  if "Kushal" in m.text and "Hegde" in m.text]
        assert person
        m = person[0]
        apply_replacement_to_block(blk, m.start_char, m.end_char, "Jane Doe")
        final = run.text
        assert "Kushal" not in final
        assert "Jane Doe" in final
        assert final.startswith("The director ")
        assert final.endswith(" signed the agreement.")


# ---------------------------------------------------------------------------
# ADDRESS bounding tests
# ---------------------------------------------------------------------------

class TestAddressBounding:
    def test_extracts_bounded_address_from_paragraph(self):
        from detectors.ner_detectors import detect_addresses
        # A 2,000+ character paragraph with an address buried in it
        prose1 = "We intend to utilize portions of the Net Proceeds for funding capital expenditure requirements for our manufacturing facilities. As of the date of this Red Herring Prospectus, we primarily manufacture our products at our four manufacturing facilities. We have recently acquired a new plot of land to expand our operations. " * 5
        address = "123, Waterloo Industrial Estate, Phase II, Bandra Kurla Complex, Mumbai 400051."
        prose2 = " The rest of the paragraph continues to describe various financial details and forward-looking statements that are entirely unrelated to the physical location of the facility. " * 5
        
        full_text = prose1 + address + prose2
        assert len(full_text) > 2000
        
        b = block(full_text)
        matches = detect_addresses(b)
        
        assert len(matches) > 0, "Failed to detect address in large paragraph"
        
        # Verify the matches are tightly bounded
        for m in matches:
            assert len(m.text) < 300, "Match is too long, bounding failed"
            # The match should exactly or closely correspond to the address string
            assert "Waterloo Industrial Estate" in m.text
            assert "Mumbai 400051" in m.text
            assert "We intend to utilize" not in m.text
            assert "The rest of the paragraph" not in m.text
