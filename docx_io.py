"""
docx_io.py
----------
DOCX text extraction and in-place run-level rewriting.

Public API
----------
load_document(path)              -> Document
extract_all_text_blocks(doc)     -> list[TextBlock]
save_document(doc, path)         -> None

Internal helpers (used by replacer.py)
---------------------------------------
_build_run_offsets(runs)         -> list[int]
apply_replacement_to_block(block, start, end, replacement) -> None
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.ns import qn

from models import TextBlock


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------

def load_document(path: str | Path) -> Document:
    """Open a .docx file and return the python-docx Document object."""
    return Document(str(path))


def save_document(doc: Document, path: str | Path) -> None:
    """Save the (modified) Document to *path*."""
    doc.save(str(path))


# ---------------------------------------------------------------------------
# Run-offset helpers
# ---------------------------------------------------------------------------

def _build_run_offsets(runs) -> list[int]:
    """
    Return a list of cumulative start-character offsets for each run.

    Example: runs with text ["Hello ", "world", "!"]
             → offsets = [0, 6, 11]
    """
    offsets: list[int] = []
    pos = 0
    for run in runs:
        offsets.append(pos)
        pos += len(run.text)
    return offsets


def _combine_run_text(runs) -> str:
    """Concatenate the text of all runs into a single string."""
    return "".join(r.text for r in runs)


# ---------------------------------------------------------------------------
# Extraction iterators
# ---------------------------------------------------------------------------

def _make_block(text: str, location: dict) -> TextBlock | None:
    """Return a TextBlock only if text is non-empty."""
    stripped = text  # preserve whitespace for offset accuracy
    if not stripped.strip():
        return None
    return TextBlock(text=text, location=location)


def iter_paragraph_blocks(doc: Document) -> Iterator[TextBlock]:
    """
    Yield one TextBlock per body paragraph that contains non-whitespace text.

    Location keys: type, para_idx, runs, run_offsets
    """
    for para_idx, para in enumerate(doc.paragraphs):
        runs = para.runs
        text = _combine_run_text(runs)
        block = _make_block(text, {
            "type": "paragraph",
            "para_idx": para_idx,
            "runs": runs,
            "run_offsets": _build_run_offsets(runs),
        })
        if block:
            yield block


def iter_table_cell_blocks(doc: Document) -> Iterator[TextBlock]:
    """
    Yield one TextBlock per non-empty paragraph in every table cell.

    Location keys: type, table_idx, row_idx, col_idx, para_idx, runs, run_offsets
    """
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    runs = para.runs
                    text = _combine_run_text(runs)
                    block = _make_block(text, {
                        "type": "table_cell",
                        "table_idx": t_idx,
                        "row_idx": r_idx,
                        "col_idx": c_idx,
                        "para_idx": p_idx,
                        "runs": runs,
                        "run_offsets": _build_run_offsets(runs),
                    })
                    if block:
                        yield block


def iter_header_footer_blocks(doc: Document) -> Iterator[TextBlock]:
    """
    Yield TextBlocks from headers and footers of all sections.

    Location keys: type ("header"|"footer"), section_idx, para_idx,
                   runs, run_offsets
    """
    for s_idx, section in enumerate(doc.sections):
        for hf_type, hf_obj in [("header", section.header),
                                 ("footer", section.footer)]:
            if hf_obj is None:
                continue
            for p_idx, para in enumerate(hf_obj.paragraphs):
                runs = para.runs
                text = _combine_run_text(runs)
                block = _make_block(text, {
                    "type": hf_type,
                    "section_idx": s_idx,
                    "para_idx": p_idx,
                    "runs": runs,
                    "run_offsets": _build_run_offsets(runs),
                })
                if block:
                    yield block


def extract_all_text_blocks(doc: Document) -> list[TextBlock]:
    """
    Extract all text from a DOCX as a flat ordered list of TextBlock objects.

    Order: body paragraphs → table cells → headers → footers.

    Each TextBlock carries the original python-docx Run objects and their
    cumulative character offsets so that replacer.py can map detected
    (start_char, end_char) spans back to the exact run(s) to modify.

    Known limitation
    ----------------
    Hyperlink text stored inside <w:hyperlink> XML elements may not be fully
    accessible via python-docx's public .runs API.  Any hyperlink text that
    is not exposed as a regular Run will be missed.  This is documented as a
    known limitation and will not cause crashes — it means hyperlinked PII
    may be under-redacted.
    """
    blocks: list[TextBlock] = []
    blocks.extend(iter_paragraph_blocks(doc))
    blocks.extend(iter_table_cell_blocks(doc))
    blocks.extend(iter_header_footer_blocks(doc))
    return blocks


# ---------------------------------------------------------------------------
# In-place run-level replacement (called by replacer.py)
# ---------------------------------------------------------------------------

def apply_replacement_to_block(
    block: TextBlock,
    start_char: int,
    end_char: int,
    replacement: str,
) -> None:
    """
    Replace the substring block.text[start_char:end_char] with *replacement*
    directly in the python-docx Run objects, preserving all run formatting.

    Strategy — single-run span
    --------------------------
    If the span falls entirely within one run, replace the substring in
    run.text and return.

    Strategy — multi-run span
    -------------------------
    When the PII span crosses one or more run boundaries (e.g. because Word
    applied bold or italic mid-entity):

    1. Identify all runs whose text overlaps [start_char, end_char).
    2. Concatenate their text to verify the original match text.
    3. Place the full *replacement* string into the FIRST affected run,
       replacing only the portion of that run's text that overlaps the span.
    4. Clear the text of every subsequent affected run (run.text = "").
       Their formatting (bold, italic, font) is preserved on the run XML node
       but the visible characters are removed.

    Ordering note
    -------------
    Callers (replacer.py::apply_to_docx) MUST process matches from the
    LAST to the FIRST (descending start_char) within each block so that
    replacing an earlier occurrence does not shift the offsets of later ones.

    Parameters
    ----------
    block       : TextBlock whose runs are to be modified in place.
    start_char  : Start offset within block.text (inclusive).
    end_char    : End offset within block.text (exclusive).
    replacement : The fake value to insert.
    """
    runs = block.location["runs"]
    offsets = block.location["run_offsets"]

    if not runs:
        return  # nothing to rewrite

    n = len(runs)

    # Find which runs overlap [start_char, end_char)
    affected: list[int] = []
    for i, run_start in enumerate(offsets):
        run_end = offsets[i + 1] if i + 1 < n else run_start + len(runs[i].text)
        # overlap condition
        if run_start < end_char and run_end > start_char:
            affected.append(i)

    if not affected:
        return  # span not found in runs (e.g. empty block edge case)

    if len(affected) == 1:
        # --- Single-run span ---
        i = affected[0]
        run = runs[i]
        run_start = offsets[i]
        local_start = start_char - run_start
        local_end = end_char - run_start
        run.text = run.text[:local_start] + replacement + run.text[local_end:]
    else:
        # --- Multi-run span ---
        first_i = affected[0]
        run = runs[first_i]
        run_start = offsets[first_i]
        local_start = start_char - run_start

        # Determine how much of the first run is part of the span
        first_run_end = offsets[first_i + 1] if first_i + 1 < n \
            else run_start + len(run.text)
        local_end_in_first = min(end_char, first_run_end) - run_start

        # Replace first run: keep prefix + insert replacement + drop suffix
        run.text = run.text[:local_start] + replacement + run.text[local_end_in_first:]

        # Clear all subsequent affected runs
        for i in affected[1:]:
            affected_run = runs[i]
            run_start_i = offsets[i]
            run_end_i = offsets[i + 1] if i + 1 < n \
                else run_start_i + len(affected_run.text)
            # How much of this run is inside the span?
            local_clear_start = max(start_char, run_start_i) - run_start_i
            local_clear_end = min(end_char, run_end_i) - run_start_i
            affected_run.text = (
                affected_run.text[:local_clear_start]
                + affected_run.text[local_clear_end:]
            )
